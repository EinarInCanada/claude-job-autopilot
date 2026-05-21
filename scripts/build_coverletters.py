"""
build_coverletters.py — Generates one cover letter .docx per resume variant.

Reads all configuration from ../config.yaml (relative to this script's location).
Reads cover letter content from ../variants/coverletters.yaml.

Usage:
    python3 scripts/build_coverletters.py

Output:
    {config.paths.output_dir}/cover_letters/  — one .docx per variant

Requirements:
    pip install python-docx lxml pyyaml
"""

import io, os, sys, yaml
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from lxml import etree

# ─── locate config ────────────────────────────────────────────────────────────

SCRIPT_DIR    = Path(__file__).resolve().parent
REPO_ROOT     = SCRIPT_DIR.parent
CONFIG_PATH   = REPO_ROOT / "config.yaml"
VARIANTS_PATH = REPO_ROOT / "variants" / "coverletters.yaml"

if not CONFIG_PATH.exists():
    sys.exit(f"ERROR: config.yaml not found at {CONFIG_PATH}\n"
             f"Copy config.example.yaml to config.yaml and fill in your details.")

with open(CONFIG_PATH) as f:
    cfg = yaml.safe_load(f)

u = cfg["user"]
OUTPUT_DIR = Path(os.path.expanduser(cfg["paths"]["output_dir"])) / cfg["paths"]["cover_letters_subdir"]
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

if not VARIANTS_PATH.exists():
    sys.exit(
        f"ERROR: Cover letter variants not found at {VARIANTS_PATH}\n"
        f"See variants/coverletters.example.yaml for the expected format."
    )

with open(VARIANTS_PATH) as f:
    cl_cfg = yaml.safe_load(f)

# ─── hyperlink helper ─────────────────────────────────────────────────────────

def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Always add Hyperlink character style (safe even if not defined — Word will pick it up)
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    r = OxmlElement('w:r')
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)
    return hyperlink

# ─── build one cover letter ───────────────────────────────────────────────────

def build_cover_letter(variant):
    filename = variant["filename"]
    out_path = OUTPUT_DIR / filename

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Pt(72)   # 1 inch
    section.bottom_margin = Pt(72)
    section.left_margin   = Pt(72)
    section.right_margin  = Pt(72)

    # ── Header: Name + contact line ───────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_run  = name_para.add_run(u["preferred_name"])
    name_run.bold      = True
    name_run.font.size = Pt(16)

    contact_para = doc.add_paragraph()
    contact_line = (
        f"{u['email']}  ·  {u['phone']}  ·  {u['city']}"
    )
    contact_run = contact_para.add_run(contact_line)
    contact_run.font.size = Pt(10)

    # Portfolio URL as hyperlink (if provided)
    if u.get("portfolio_url"):
        contact_para.add_run("  ·  ")
        add_hyperlink(contact_para, u["portfolio_url"], u["portfolio_url"])

    doc.add_paragraph()  # spacer

    # ── Date ──────────────────────────────────────────────────────────────────
    from datetime import date
    doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    doc.add_paragraph()  # spacer

    # ── Salutation ────────────────────────────────────────────────────────────
    doc.add_paragraph(variant.get("salutation", "Dear Hiring Manager,"))
    doc.add_paragraph()

    # ── Body paragraphs ───────────────────────────────────────────────────────
    for para_text in variant["body"]:
        p = doc.add_paragraph(para_text)
        p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()

    # ── Closing ───────────────────────────────────────────────────────────────
    doc.add_paragraph(variant.get("closing", "Sincerely,"))
    doc.add_paragraph(u["preferred_name"])

    # ── Save ──────────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    with open(out_path, 'wb') as f:
        f.write(buf.read())
    print(f"  ✓  {filename}  ({os.path.getsize(out_path):,} bytes)")

# ─── run ──────────────────────────────────────────────────────────────────────

variants = cl_cfg.get("variants", [])
print(f"\nBuilding {len(variants)} cover letter(s)…")
for v in variants:
    build_cover_letter(v)

print(f"\nAll cover letters saved to: {OUTPUT_DIR}")
