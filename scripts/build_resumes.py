"""
build_resumes.py — Generates tailored resume variants from a base .docx template.

Reads all configuration from ../config.yaml (relative to this script's location).
Reads variant definitions from ../variants/resumes.yaml.

Usage:
    python3 scripts/build_resumes.py

Output:
    {config.paths.output_dir}/resumes/  — one .docx per variant defined in variants/resumes.yaml

Requirements:
    pip install python-docx lxml pyyaml
"""

import io, os, copy, sys, yaml
from lxml import etree
from docx import Document
from pathlib import Path

# ─── locate config ────────────────────────────────────────────────────────────

SCRIPT_DIR  = Path(__file__).resolve().parent
REPO_ROOT   = SCRIPT_DIR.parent
CONFIG_PATH = REPO_ROOT / "config.yaml"
VARIANTS_PATH = REPO_ROOT / "variants" / "resumes.yaml"

if not CONFIG_PATH.exists():
    sys.exit(f"ERROR: config.yaml not found at {CONFIG_PATH}\n"
             f"Copy config.example.yaml to config.yaml and fill in your details.")

with open(CONFIG_PATH) as f:
    cfg = yaml.safe_load(f)

# ─── resolve paths ────────────────────────────────────────────────────────────

OUTPUT_DIR   = Path(os.path.expanduser(cfg["paths"]["output_dir"])) / cfg["paths"]["resumes_subdir"]
TEMPLATE     = Path(os.path.expanduser(cfg["paths"]["resume_template"]))

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

if not TEMPLATE.exists():
    sys.exit(f"ERROR: Resume template not found at {TEMPLATE}\n"
             f"Update 'paths.resume_template' in config.yaml.")

# ─── namespace helpers ────────────────────────────────────────────────────────

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W  = lambda tag: f'{{{NS}}}{tag}'

def all_texts(el):
    """Return concatenated text of all w:t descendants."""
    return ''.join(t.text or '' for t in el.iter(W('t')))

def set_first_text(el, new_text):
    """Replace text in first w:t, clear the rest."""
    ts = list(el.iter(W('t')))
    if not ts:
        return
    ts[0].text = new_text
    ts[0].attrib.pop('{http://www.w3.org/XML/1998/namespace}space', None)
    if ' ' in new_text or new_text.startswith(' ') or new_text.endswith(' '):
        ts[0].set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    for t in ts[1:]:
        t.text = ''

def make_bullet_para(ref_para_el, text):
    """Clone a list-paragraph element and replace its text."""
    new_el = copy.deepcopy(ref_para_el)
    set_first_text(new_el, text)
    runs = new_el.findall(f'.//{W("r")}')
    if len(runs) > 1:
        for r in runs[1:]:
            r.getparent().remove(r)
    return new_el

def make_header_table(ref_tbl_el, title, date):
    """Clone a project header table and replace title & date."""
    new_tbl = copy.deepcopy(ref_tbl_el)
    rows = new_tbl.findall(f'.//{W("tr")}')
    if rows:
        row = rows[0]
        cells = row.findall(f'.//{W("tc")}')
        if len(cells) >= 2:
            set_first_text(cells[0], title)
            set_first_text(cells[1], date)
    return new_tbl

# ─── document helpers ─────────────────────────────────────────────────────────

def body_children(body):
    return [c for c in body if c.tag in (W('p'), W('tbl'))]

def find_tbl_by_text(body, keyword):
    for c in body_children(body):
        if c.tag == W('tbl') and keyword.lower() in all_texts(c).lower():
            return c
    return None

def find_list_paras_after(body, tbl_el):
    """Collect consecutive ListParagraph paragraphs immediately after tbl_el."""
    result, found = [], False
    for c in body_children(body):
        if found:
            if c.tag == W('p'):
                style = c.find(f'.//{W("pStyle")}')
                if style is not None and 'List' in (style.get(W('val'), '') or ''):
                    result.append(c)
                else:
                    break
            else:
                break
        if c is tbl_el:
            found = True
    return result

def replace_para_texts(para_els, new_texts):
    """
    Replace paragraph texts with new_texts.
    If new_texts is None or shorter, remove surplus paragraphs from the body entirely
    (avoids invisible empty lines that inflate page count).
    """
    if new_texts is None:
        for p in para_els:
            p.getparent().remove(p)
        return
    for i, p in enumerate(para_els):
        if i < len(new_texts):
            set_first_text(p, new_texts[i])
            runs = list(p.iter(W('r')))
            for r in runs[1:]:
                for t in r.iter(W('t')):
                    t.text = ''
        else:
            p.getparent().remove(p)

# ─── load variant definitions ─────────────────────────────────────────────────

if not VARIANTS_PATH.exists():
    sys.exit(
        f"ERROR: Variant definitions not found at {VARIANTS_PATH}\n"
        f"See variants/resumes.example.yaml for the expected format."
    )

with open(VARIANTS_PATH) as f:
    variants_cfg = yaml.safe_load(f)

VARIANTS = variants_cfg.get("variants", [])
PROJECTS = variants_cfg.get("projects", {})

# ─── build one resume ────────────────────────────────────────────────────────

def build_resume(variant):
    filename = variant["filename"]
    out_path = OUTPUT_DIR / filename

    # Fresh copy from template via BytesIO (avoids FUSE overwrite permission errors)
    with open(TEMPLATE, 'rb') as f:
        raw = f.read()

    doc = Document(io.BytesIO(raw))
    body = doc.element.body

    # ── 1. Self-Assessment / Summary ─────────────────────────────────────────
    # Assumes summary is in the first cell of the first table in the document.
    if doc.tables:
        cell = doc.tables[0].rows[0].cells[0]
        for para in cell.paragraphs:
            for run in para.runs:
                if run.text.strip():
                    run.text = variant["summary"]
                    break

    # ── 2. Update project bullet points ──────────────────────────────────────
    # Each entry in variant["projects"] is:
    #   { "find_keyword": "...", "bullets": [...] or null }
    for proj in variant.get("projects", []):
        tbl = find_tbl_by_text(body, proj["find_keyword"])
        if tbl is not None:
            paras = find_list_paras_after(body, tbl)
            replace_para_texts(paras, proj.get("bullets"))

    # ── 3. Insert new projects (newest first) ────────────────────────────────
    # Each entry in variant["insert_projects"] describes a project to insert
    # BEFORE a reference project already in the document.
    # Format:
    #   { "title": "...", "date": "...", "bullets": [...],
    #     "insert_before_keyword": "...", "clone_table_keyword": "..." }
    ref_bullet_cache = {}

    for proj in variant.get("insert_projects", []):
        anchor_tbl = find_tbl_by_text(body, proj["insert_before_keyword"])
        clone_src  = find_tbl_by_text(body, proj["clone_table_keyword"])
        if anchor_tbl is None or clone_src is None:
            print(f"  ⚠  Could not insert project '{proj['title']}' — anchor or clone table not found")
            continue

        new_tbl = make_header_table(clone_src, proj["title"], proj["date"])
        anchor_idx = list(body).index(anchor_tbl)
        body.insert(anchor_idx, new_tbl)

        # Use first existing bullet after clone_src as reference paragraph
        if proj["clone_table_keyword"] not in ref_bullet_cache:
            ref_bullets = find_list_paras_after(body, clone_src)
            ref_bullet_cache[proj["clone_table_keyword"]] = ref_bullets[0] if ref_bullets else None

        ref_para = ref_bullet_cache[proj["clone_table_keyword"]]
        if ref_para is not None:
            new_tbl_idx = list(body).index(new_tbl)
            for i, text in enumerate(proj.get("bullets", [])):
                bullet = make_bullet_para(ref_para, text)
                body.insert(new_tbl_idx + 1 + i, bullet)

    # ── 4. Update Skills section ──────────────────────────────────────────────
    # Matches paragraphs whose text starts with any known skill-row prefix.
    skill_prefixes = tuple(variant.get("skill_prefixes", [
        "Programming Languages:", "ML &", "Frameworks &", "Tools:", "Data &",
        "IT &", "Concepts:", "PM Tools:", "AI Dev Tools:", "AI Workflow",
        "Technical Skills:", "AI &", "Concepts &"
    ]))
    skill_paras = [p for p in doc.paragraphs
                   if p.text.strip().startswith(skill_prefixes)]
    for i, para in enumerate(skill_paras):
        if i < len(variant["skills"]):
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = variant["skills"][i]
            else:
                para.add_run(variant["skills"][i])

    # ── 5. Save ───────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    with open(out_path, 'wb') as f:
        f.write(buf.read())
    print(f"  ✓  {filename}  ({os.path.getsize(out_path):,} bytes)")

# ─── run ──────────────────────────────────────────────────────────────────────

print(f"\nBuilding {len(VARIANTS)} resume variant(s)…")
for v in VARIANTS:
    build_resume(v)

print(f"\nAll resumes saved to: {OUTPUT_DIR}")
print("\nNext step: convert to PDF with LibreOffice:")
print(f"  soffice --headless --convert-to pdf {OUTPUT_DIR}/*.docx --outdir {OUTPUT_DIR}")
